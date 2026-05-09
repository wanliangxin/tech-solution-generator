"""
文档解析服务 — 支持 PDF 和 DOCX 格式
负责从技术规范书中提取结构化目录（TOC）及各章节原始内容。

解析策略（TOC 优先）：
  Phase 1: 提取目录结构（Word SDT TOC → 手动目录 → 回退标题扫描）
  Phase 2: 将目录条目映射到正文中对应位置，提取区间内容
  Phase 3: 对每个区间内容做规则式提炼，生成 content_hint

层级识别规则：
  level 1 — 一、二、三、…（中文序数 + 顿号）/ 第X章
  level 2 — 1.  2.  3.  …（阿拉伯数字 + 点 + 空格）
  level 3 — （1）（2）（3）…（全角/半角括号包裹数字）
  level 4 — ① ② ③ … 或 1.1  1.2  2.1 …（带圈数字 / 小数点编号）

特殊标记：
  ★ — 否决条款（实质性条款），保留在 special_marks 中
  ▲ — 加分项，保留在 special_marks 中
"""

import os
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Callable, Union, BinaryIO
import logging

logger = logging.getLogger(__name__)

CONTENT_HINT_LIMIT = 2000

_CIRCLED_DIGITS = set("①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳")

# TOC 样式名（Word 内置 TOC 段落样式）
_TOC_STYLES = {
    "toc 1", "toc 2", "toc 3", "toc 4",
    "toc1", "toc2", "toc3", "toc4",
    "目录 1", "目录 2", "目录 3", "目录 4",
}

# 手动目录区域的标记词
_MANUAL_TOC_MARKERS = {"目录", "目 录", "contents", "table of contents"}

# 需求关键词（用于内容提炼）
_REQUIREMENT_KEYWORDS = re.compile(
    r'(要求|必须|应当|应|需要|需|不低于|不少于|不超过|不得|禁止|至少|满足|符合|达到|具备|保证|确保)'
)


# ─────────────────────────────────────────────
# 数据结构
# ─────────────────────────────────────────────

class Section:
    """代表文档中的一个章节"""
    def __init__(
        self,
        section_id: str,
        level: int,
        title: str,
        content_hint: str = "",
        special_marks: Optional[list] = None,
    ):
        self.id = section_id
        self.level = level
        self.title = title
        self.content_hint = content_hint[:CONTENT_HINT_LIMIT]
        self.special_marks: list[str] = special_marks or []

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "level": self.level,
            "title": self.title,
            "content_hint": self.content_hint,
            "special_marks": self.special_marks,
        }


class ParsedDocument:
    """解析后的文档结构"""
    def __init__(self, doc_id: str, title: str, sections: list[Section]):
        self.doc_id = doc_id
        self.title = title
        self.sections = sections

    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "title": self.title,
            "sections": [s.to_dict() for s in self.sections],
        }


@dataclass
class TocEntry:
    """目录中的一个条目"""
    level: int
    title: str
    raw_title: str
    special_marks: list[str] = field(default_factory=list)


# ─────────────────────────────────────────────
# 特殊标记提取
# ─────────────────────────────────────────────

_SPECIAL_MARK_RE = re.compile(r'[★▲]')


def extract_special_marks(text: str) -> tuple[list[str], str]:
    marks = []
    for m in ("★", "▲"):
        if m in text:
            marks.append(m)
    cleaned = _SPECIAL_MARK_RE.sub("", text).strip()
    return marks, cleaned


# ─────────────────────────────────────────────
# 编号识别正则（4层级）
# ─────────────────────────────────────────────

HEADING_PATTERNS: list[tuple[int, re.Pattern]] = [
    (4, re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]')),
    (4, re.compile(r'^\d+\.\d+[\s　]')),
    (3, re.compile(r'^（\d+）')),
    (3, re.compile(r'^\(\d+\)')),
    (2, re.compile(r'^\d+\.\s')),
    (2, re.compile(r'^\d+\.　')),
    (1, re.compile(r'^[一二三四五六七八九十百千]+[、，]')),
    (1, re.compile(r'^第\s*[一二三四五六七八九十百千\d]+\s*[章节篇部]')),
]


def detect_level_from_numbering(text: str) -> Optional[int]:
    text = text.strip()
    for level, pattern in HEADING_PATTERNS:
        if pattern.match(text):
            return level
    return None


def clean_title(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


# ─────────────────────────────────────────────
# TOC 提取：Word 内置 SDT
# ─────────────────────────────────────────────

def _extract_toc_from_sdt(document) -> list[TocEntry]:
    """从 Word 内置 TOC（SDT 结构）中提取目录条目"""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        return []

    body = document.element.body
    entries: list[TocEntry] = []

    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find(qn("w:sdtPr"))
        if sdt_pr is None:
            continue
        doc_part_obj = sdt_pr.find(qn("w:docPartObj"))
        if doc_part_obj is not None:
            gallery = doc_part_obj.find(qn("w:docPartGallery"))
            if gallery is not None and "Table of Contents" in (gallery.get(qn("w:val")) or ""):
                sdt_content = sdt.find(qn("w:sdtContent"))
                if sdt_content is not None:
                    entries = _parse_toc_paragraphs_from_xml(sdt_content, document)
                    if entries:
                        return entries

    # 备用：查找 body 级别的 TOC 样式段落（不在 SDT 内的情况）
    from docx.text.paragraph import Paragraph
    for p_elem in body.iterchildren(qn("w:p")):
        para = Paragraph(p_elem, document._body)
        style_name = (para.style.name or "").lower().strip() if para.style else ""
        if style_name in _TOC_STYLES:
            text = para.text.strip()
            if text:
                level = _toc_style_to_level(style_name)
                text = _clean_toc_line(text)
                marks, cleaned = extract_special_marks(text)
                if cleaned:
                    entries.append(TocEntry(level=level, title=cleaned, raw_title=text, special_marks=marks))

    return entries


def _find_sdt_toc_end(paragraphs) -> int:
    """
    扫描 paragraphs，找到最后一个 TOC 样式段落的索引+1。
    用于 SDT 目录路径的 body_start，确保正文匹配从目录区之后开始。
    """
    last_toc_idx = -1
    scan_limit = min(len(paragraphs), 300)
    for i in range(scan_limit):
        style_name = (paragraphs[i].style.name or "").lower().strip() if paragraphs[i].style else ""
        if style_name in _TOC_STYLES:
            last_toc_idx = i
    return last_toc_idx + 1 if last_toc_idx >= 0 else 0


def _parse_toc_paragraphs_from_xml(sdt_content, document) -> list[TocEntry]:
    """从 SDT Content 的 XML 中解析 TOC 段落"""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    entries: list[TocEntry] = []
    for p_elem in sdt_content.iter(qn("w:p")):
        para = Paragraph(p_elem, document._body)
        style_name = (para.style.name or "").lower().strip() if para.style else ""
        if style_name not in _TOC_STYLES:
            continue
        text = para.text.strip()
        if not text:
            continue
        level = _toc_style_to_level(style_name)
        text = _clean_toc_line(text)
        marks, cleaned = extract_special_marks(text)
        if cleaned:
            entries.append(TocEntry(level=level, title=cleaned, raw_title=text, special_marks=marks))
    return entries


def _toc_style_to_level(style_name: str) -> int:
    """TOC 样式名转层级数字"""
    for i in range(4, 0, -1):
        if str(i) in style_name:
            return i
    return 1


# ─────────────────────────────────────────────
# TOC 提取：手动编写的目录
# ─────────────────────────────────────────────

def _extract_toc_manual(paragraphs) -> tuple[list[TocEntry], int]:
    """扫描文档前部，查找手动编写的目录区域并提取条目。
    返回 (entries, toc_end_idx)，toc_end_idx 为目录区最后一段的索引+1，
    供调用方跳过目录区直接在正文中匹配标题。
    """
    toc_start_idx = -1
    scan_limit = min(len(paragraphs), 80)

    for i in range(scan_limit):
        text = paragraphs[i].text.strip().lower()
        if text in _MANUAL_TOC_MARKERS:
            toc_start_idx = i + 1
            break

    if toc_start_idx < 0:
        return [], 0

    entries: list[TocEntry] = []
    blank_count = 0
    toc_end_idx = toc_start_idx

    for i in range(toc_start_idx, min(len(paragraphs), toc_start_idx + 200)):
        raw_text = paragraphs[i].text.strip()

        if not raw_text:
            blank_count += 1
            if blank_count >= 2 and entries:
                break
            continue
        blank_count = 0

        # 目录区域结束判定：出现长正文段落
        if len(raw_text) > 120 and not detect_level_from_numbering(raw_text):
            break

        cleaned = _clean_toc_line(raw_text)
        if not cleaned:
            continue

        marks, title = extract_special_marks(cleaned)
        level = detect_level_from_numbering(title)

        if level is not None and len(title) < 100:
            entries.append(TocEntry(level=level, title=title, raw_title=raw_text, special_marks=marks))
            toc_end_idx = i + 1
        elif entries and len(title) < 80:
            # 可能是无编号的顶级标题（如 "技术方案概述"）
            style_name = (paragraphs[i].style.name or "").lower() if paragraphs[i].style else ""
            if "heading" in style_name or "标题" in style_name:
                entries.append(TocEntry(level=1, title=title, raw_title=raw_text, special_marks=marks))
                toc_end_idx = i + 1

    return entries, toc_end_idx


def _clean_toc_line(text: str) -> str:
    """清理 TOC 行：去除尾部导引符、页码、Tab"""
    text = re.sub(r'[\t]+', ' ', text)
    text = re.sub(r'[.·…]{3,}\s*\d*\s*$', '', text)
    text = re.sub(r'\s+\d+\s*$', '', text)
    return text.strip()


# ─────────────────────────────────────────────
# 内容映射：TOC 条目 → 正文位置
# ─────────────────────────────────────────────

def _map_toc_to_body(
    toc_entries: list[TocEntry],
    paragraphs,
    progress_callback: Optional[Callable] = None,
    body_start: int = 0,
) -> list[tuple[TocEntry, str]]:
    """
    将目录条目映射到正文中对应的标题位置，提取该标题到下一个标题之间的内容。
    body_start: 从此段落索引开始搜索正文标题，用于跳过文档前部的目录区域。
    """
    if progress_callback:
        progress_callback("映射目录到正文...", 0, len(toc_entries))

    # 构建段落文本索引
    para_texts = [(i, paragraphs[i].text.strip()) for i in range(len(paragraphs))]

    # 为每个 TOC 条目找到正文中对应的段落位置
    matched_positions: list[int] = []
    last_pos = body_start

    for entry in toc_entries:
        pos = _find_heading_in_body(entry, para_texts, last_pos)
        matched_positions.append(pos)
        if pos >= 0:
            last_pos = pos + 1

    matched_count = sum(1 for p in matched_positions if p >= 0)
    logger.info(
        f"TOC正文匹配结果：{matched_count}/{len(toc_entries)} 条命中，"
        f"body_start={body_start}，"
        f"命中位置范围：{[p for p in matched_positions if p >= 0][:5]}..."
    )

    # 提取每个条目对应区间的正文
    results: list[tuple[TocEntry, str]] = []
    for idx, (entry, pos) in enumerate(zip(toc_entries, matched_positions)):
        if progress_callback:
            progress_callback(entry.title, idx + 1, len(toc_entries))

        if pos < 0:
            results.append((entry, ""))
            continue

        # 找下一个已匹配的标题位置作为区间终点
        next_pos = len(paragraphs)
        for j in range(idx + 1, len(matched_positions)):
            if matched_positions[j] >= 0:
                next_pos = matched_positions[j]
                break

        # 提取区间内的正文段落
        content_parts = []
        for j in range(pos + 1, next_pos):
            text = paragraphs[j].text.strip()
            if text:
                content_parts.append(text)

        raw_content = "\n".join(content_parts)
        results.append((entry, raw_content))

    return results


def _find_heading_in_body(entry: TocEntry, para_texts: list[tuple[int, str]], start_from: int) -> int:
    """在段落列表中找到与 TOC 条目匹配的标题位置"""
    target = _normalize_for_match(entry.title)
    if not target:
        return -1

    # 提取编号前缀用于精确匹配
    target_prefix = _extract_numbering_prefix(entry.title)

    best_idx = -1
    best_score = 0

    for i, text in para_texts:
        if i < start_from:
            continue
        if not text or len(text) > 200:
            continue

        normalized = _normalize_for_match(text)

        # 策略1：编号前缀 + 文本前部匹配
        if target_prefix:
            text_prefix = _extract_numbering_prefix(text)
            if text_prefix and text_prefix == target_prefix:
                score = _text_similarity(target, normalized)
                if score > best_score:
                    best_score = score
                    best_idx = i
                    if score > 0.8:
                        return i

        # 策略2：文本高度相似
        score = _text_similarity(target, normalized)
        if score > 0.85 and score > best_score:
            best_score = score
            best_idx = i

    return best_idx


def _normalize_for_match(text: str) -> str:
    """归一化文本用于匹配：去除标记、空白、标点"""
    text = _SPECIAL_MARK_RE.sub("", text)
    text = re.sub(r'[.·…]+\s*\d*\s*$', '', text)
    text = re.sub(r'\s+', '', text)
    return text.lower()


def _extract_numbering_prefix(text: str) -> str:
    """提取编号前缀（如 "一、"、"1."、"（1）"）"""
    text = text.strip()
    for _, pattern in HEADING_PATTERNS:
        m = pattern.match(text)
        if m:
            return m.group(0)
    return ""


def _text_similarity(a: str, b: str) -> float:
    """简单的文本相似度（基于公共前缀长度比）"""
    if not a or not b:
        return 0.0
    shorter = min(len(a), len(b))
    common = 0
    for i in range(shorter):
        if a[i] == b[i]:
            common += 1
        else:
            break
    prefix_ratio = common / shorter if shorter > 0 else 0

    # 也考虑包含关系
    if a in b or b in a:
        return 0.9

    return prefix_ratio


# ─────────────────────────────────────────────
# 内容提炼（规则式，不调 LLM）
# ─────────────────────────────────────────────

def _distill_content(raw_content: str, max_chars: int = CONTENT_HINT_LIMIT) -> str:
    """
    从原始正文中提炼核心内容。
    短内容原样保留，长内容提取：开头句 + 需求句 + 量化指标句。
    """
    if not raw_content:
        return ""
    if len(raw_content) <= max_chars:
        return raw_content

    sentences = _split_sentences(raw_content)
    if not sentences:
        return raw_content[:max_chars]

    parts: list[str] = []
    used_indices: set[int] = set()

    # 1. 开头 3-5 句作为上下文
    opening_budget = min(5, len(sentences))
    for i in range(opening_budget):
        parts.append(sentences[i])
        used_indices.add(i)
        if sum(len(p) for p in parts) > max_chars // 3:
            break

    # 2. 含需求关键词的句子
    for i, sent in enumerate(sentences):
        if i in used_indices:
            continue
        if _REQUIREMENT_KEYWORDS.search(sent):
            parts.append(sent)
            used_indices.add(i)
            if sum(len(p) for p in parts) > max_chars * 0.8:
                break

    # 3. 含量化指标的句子（数字+单位）
    quant_re = re.compile(r'\d+\s*[%％倍天小时分钟秒个台套件项万亿元GB|TB|MB|Gbps|Mbps]')
    for i, sent in enumerate(sentences):
        if i in used_indices:
            continue
        if quant_re.search(sent):
            parts.append(sent)
            used_indices.add(i)
            if sum(len(p) for p in parts) > max_chars * 0.9:
                break

    result = "\n".join(parts)
    return result[:max_chars]


def _split_sentences(text: str) -> list[str]:
    """按中文/英文句子分割"""
    parts = re.split(r'(?<=[。；;！？\n])', text)
    sentences = [p.strip() for p in parts if p.strip()]
    return sentences


# ─────────────────────────────────────────────
# DOCX 解析（TOC 优先策略）
# ─────────────────────────────────────────────

def parse_docx(
    file_source: Union[str, BinaryIO],
    filename: str = "",
    progress_callback: Optional[Callable[[str, int, int], None]] = None,
) -> ParsedDocument:
    """
    解析 DOCX 文档。策略：
    1. 优先从文档中提取目录结构（Word SDT TOC 或手动目录）
    2. 将目录条目映射到正文，提取对应区间内容
    3. 对区间内容做提炼后作为 content_hint
    4. 无法找到目录时回退到标题扫描
    """
    from docx import Document

    doc_id = str(uuid.uuid4())

    if progress_callback:
        progress_callback("读取文档...", 0, -1)

    document = Document(file_source)
    paragraphs = document.paragraphs

    # 提取文档标题
    doc_title = Path(filename).stem if filename else "文档"
    for para in paragraphs[:20]:
        style_name = (para.style.name or "").lower() if para.style else ""
        if "title" in style_name and para.text.strip():
            doc_title = clean_title(para.text)
            break

    # Phase 1: 提取 TOC
    if progress_callback:
        progress_callback("提取目录结构...", 0, -1)

    toc_entries = _extract_toc_from_sdt(document)
    toc_source = "Word内置目录"
    body_start = 0

    if toc_entries:
        # SDT TOC 的条目段落被展开到 paragraphs 中，需找到 TOC 样式段落的结束位置
        # 以此作为 body_start，避免 _map_toc_to_body 命中 TOC 区块本身
        body_start = _find_sdt_toc_end(paragraphs)

    if not toc_entries:
        toc_entries, body_start = _extract_toc_manual(paragraphs)
        toc_source = "手动目录"

    logger.info(f"TOC提取路径：{toc_source}，条目数：{len(toc_entries)}，body_start：{body_start}")

    if not toc_entries:
        if progress_callback:
            progress_callback("未发现目录，使用标题扫描...", 0, -1)
        return _parse_docx_by_heading_scan(document, doc_id, doc_title, progress_callback)

    logger.info(f"从{toc_source}提取到 {len(toc_entries)} 个章节")
    if progress_callback:
        progress_callback(f"从{toc_source}提取到 {len(toc_entries)} 个章节", 0, len(toc_entries))

    # Phase 2: 映射到正文（body_start 跳过目录区，避免命中目录页本身）
    mapped = _map_toc_to_body(toc_entries, paragraphs, progress_callback, body_start=body_start)

    # Phase 3: 构建 Section 列表（含内容提炼）
    if progress_callback:
        progress_callback("提炼章节内容...", 0, len(mapped))

    sections: list[Section] = []
    for idx, (entry, raw_content) in enumerate(mapped):
        distilled = _distill_content(raw_content)
        sections.append(Section(
            section_id=f"s{idx + 1}",
            level=entry.level,
            title=clean_title(entry.title),
            content_hint=distilled,
            special_marks=entry.special_marks,
        ))
        if progress_callback:
            progress_callback(entry.title, idx + 1, len(mapped))

    return ParsedDocument(doc_id, doc_title, sections)


# ─────────────────────────────────────────────
# 回退：标题扫描（原有逻辑）
# ─────────────────────────────────────────────

def _parse_docx_by_heading_scan(
    document,
    doc_id: str,
    doc_title: str,
    progress_callback: Optional[Callable] = None,
) -> ParsedDocument:
    """当无法提取 TOC 时，回退到逐段落扫描标题的方式"""
    paragraphs = document.paragraphs
    sections: list[Section] = []

    if progress_callback:
        progress_callback("扫描章节标题...", 0, -1)

    heading_indices: list[tuple[int, int, str, list[str]]] = []
    for i, para in enumerate(paragraphs):
        raw_text = para.text.strip()
        if not raw_text:
            continue
        marks, text = extract_special_marks(raw_text)
        style_name = para.style.name if para.style else ""
        level = _get_heading_level(style_name, text)
        if level is not None:
            heading_indices.append((i, level, text, marks))

    total_headings = len(heading_indices)
    if progress_callback:
        progress_callback(f"发现 {total_headings} 个章节", 0, total_headings)

    for h_idx, (para_idx, level, text, marks) in enumerate(heading_indices):
        if progress_callback:
            progress_callback(text, h_idx + 1, total_headings)

        next_para_idx = (
            heading_indices[h_idx + 1][0]
            if h_idx + 1 < len(heading_indices)
            else len(paragraphs)
        )

        hint_parts = []
        for j in range(para_idx + 1, next_para_idx):
            next_text = paragraphs[j].text.strip()
            if next_text:
                hint_parts.append(next_text)

        raw_content = "\n".join(hint_parts)
        distilled = _distill_content(raw_content)

        sections.append(Section(
            section_id=f"s{h_idx + 1}",
            level=level,
            title=clean_title(text),
            content_hint=distilled,
            special_marks=marks,
        ))

    return ParsedDocument(doc_id, doc_title, sections)


def _get_heading_level(style_name: str, text: str) -> Optional[int]:
    s = style_name.lower().strip()
    heading_map = {
        "heading 1": 1, "heading1": 1, "标题 1": 1, "标题1": 1,
        "heading 2": 2, "heading2": 2, "标题 2": 2, "标题2": 2,
        "heading 3": 3, "heading3": 3, "标题 3": 3, "标题3": 3,
        "heading 4": 4, "heading4": 4, "标题 4": 4, "标题4": 4,
        "heading 5": 4, "heading5": 4, "标题 5": 4, "标题5": 4,
        "heading 6": 4, "heading6": 4, "标题 6": 4, "标题6": 4,
    }
    for key, level in heading_map.items():
        if s == key.lower():
            return level

    if len(text) < 120:
        numbering_level = detect_level_from_numbering(text)
        if numbering_level is not None:
            return numbering_level

    return None


# ─────────────────────────────────────────────
# PDF 解析
# ─────────────────────────────────────────────

def parse_pdf(
    file_source: Union[str, BinaryIO],
    filename: str = "",
    progress_callback: Optional[Callable[[str, int, int], None]] = None,
) -> ParsedDocument:
    """解析 PDF 文档，提取目录结构与各章节原始内容。"""
    import pdfplumber

    doc_id = str(uuid.uuid4())
    sections: list[Section] = []
    doc_title = filename or "PDF Document"

    with pdfplumber.open(file_source) as pdf:
        total_pages = len(pdf.pages)
        if progress_callback:
            progress_callback("读取 PDF 页面...", 0, total_pages)

        all_words: list[dict] = []
        for page_no, page in enumerate(pdf.pages):
            words = page.extract_words(
                extra_attrs=["fontname", "size"],
                use_text_flow=True,
            )
            all_words.extend(words)
            if progress_callback and page_no % 5 == 0:
                progress_callback(f"扫描第 {page_no+1}/{total_pages} 页", page_no + 1, total_pages)

        if not all_words:
            logger.warning("PDF 未能提取到文字，可能是扫描件")
            return ParsedDocument(doc_id, doc_title, [])

        sizes = [w.get("size", 0) for w in all_words if w.get("size", 0) > 0]
        if not sizes:
            return ParsedDocument(doc_id, doc_title, [])

        from collections import Counter
        size_counter = Counter(round(s, 1) for s in sizes)
        body_size = size_counter.most_common(1)[0][0]

        lines = _group_words_into_lines(all_words)

        if progress_callback:
            progress_callback("识别章节标题...", 0, -1)

        heading_indices: list[tuple[int, int, str, list[str]]] = []
        for i, line in enumerate(lines):
            raw_text = line["text"].strip()
            if not raw_text:
                continue

            marks, text = extract_special_marks(raw_text)
            avg_size = line["avg_size"]
            numbering_level = detect_level_from_numbering(text)

            is_heading = False
            level = 1

            if numbering_level is not None:
                is_heading = True
                level = numbering_level
            elif avg_size > body_size * 1.15:
                is_heading = True
                level = 1
            elif avg_size > body_size * 1.05 and len(text) < 60:
                is_heading = True
                level = 2

            if is_heading and len(text) >= 2:
                heading_indices.append((i, level, text, marks))

        total_headings = len(heading_indices)
        for h_idx, (line_idx, level, text, marks) in enumerate(heading_indices):
            if progress_callback:
                progress_callback(text, h_idx + 1, total_headings)

            next_line_idx = (
                heading_indices[h_idx + 1][0]
                if h_idx + 1 < len(heading_indices)
                else len(lines)
            )
            hint_parts = []
            for j in range(line_idx + 1, next_line_idx):
                next_text = lines[j]["text"].strip()
                if next_text:
                    hint_parts.append(next_text)
            raw_content = "\n".join(hint_parts)
            distilled = _distill_content(raw_content)

            if h_idx == 0 and level == 1:
                doc_title = clean_title(text)

            sections.append(Section(
                section_id=f"s{h_idx + 1}",
                level=level,
                title=clean_title(text),
                content_hint=distilled,
                special_marks=marks,
            ))

    return ParsedDocument(doc_id, doc_title, sections)


def _group_words_into_lines(words: list[dict]) -> list[dict]:
    if not words:
        return []
    lines = []
    current_line_words = [words[0]]
    for word in words[1:]:
        prev = current_line_words[-1]
        y_diff = abs(word.get("top", 0) - prev.get("top", 0))
        avg_size = prev.get("size", 12) or 12
        if y_diff < avg_size * 0.5:
            current_line_words.append(word)
        else:
            lines.append(_words_to_line(current_line_words))
            current_line_words = [word]
    lines.append(_words_to_line(current_line_words))
    return lines


def _words_to_line(words: list[dict]) -> dict:
    text = " ".join(w.get("text", "") for w in words)
    sizes = [w.get("size", 0) for w in words if w.get("size", 0) > 0]
    avg_size = sum(sizes) / len(sizes) if sizes else 12
    return {"text": text, "avg_size": avg_size}


# ─────────────────────────────────────────────
# 统一入口
# ─────────────────────────────────────────────

def parse_document(
    file_source: Union[str, BinaryIO],
    suffix: str = "",
    filename: str = "",
    progress_callback: Optional[Callable[[str, int, int], None]] = None,
) -> ParsedDocument:
    if isinstance(file_source, str):
        suffix = suffix or Path(file_source).suffix.lower()
        filename = filename or Path(file_source).stem

    suffix = suffix.lower()

    actual_type = _detect_file_type(file_source)
    if actual_type == "ole" and suffix == ".docx":
        logger.warning(f"文件扩展名为 .docx 但实际是 OLE (.doc) 格式：{filename}")
        suffix = ".doc"

    if suffix == ".pdf":
        return parse_pdf(file_source, filename, progress_callback)
    elif suffix == ".docx":
        return parse_docx(file_source, filename, progress_callback)
    elif suffix == ".doc":
        return _parse_doc_via_convert(file_source, filename, progress_callback)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}，请上传 PDF 或 DOCX 文件")


def _detect_file_type(file_source: Union[str, BinaryIO]) -> str:
    if isinstance(file_source, str):
        try:
            with open(file_source, "rb") as f:
                header = f.read(8)
        except Exception:
            return "unknown"
    else:
        pos = file_source.tell()
        header = file_source.read(8)
        file_source.seek(pos)

    if header[:4] == b'PK\x03\x04':
        return "zip"
    if header[:4] == b'\xd0\xcf\x11\xe0':
        return "ole"
    if header[:5] == b'%PDF-':
        return "pdf"
    return "unknown"


def _parse_doc_via_convert(
    file_source: Union[str, BinaryIO],
    filename: str = "",
    progress_callback: Optional[Callable[[str, int, int], None]] = None,
) -> ParsedDocument:
    import subprocess
    import tempfile

    if progress_callback:
        progress_callback("检测到旧版 .doc 格式，正在转换...", 0, -1)

    out_dir = tempfile.mkdtemp(prefix="tsg_doc_convert_")
    stem = Path(filename).stem if filename else "document"

    if isinstance(file_source, str):
        doc_path = file_source
    else:
        doc_path = os.path.join(out_dir, f"{stem}.doc")
        with open(doc_path, "wb") as f:
            f.write(file_source.read())

    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", out_dir],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise ValueError(f".doc 转换失败：{result.stderr.strip()}")
    except FileNotFoundError:
        raise ValueError(
            "该文件实际为旧版 .doc 格式（非 .docx），"
            "服务器未安装 LibreOffice 无法自动转换。"
            "请先用 Word 或 WPS 将文件「另存为」.docx 格式后重新上传。"
        )
    except subprocess.TimeoutExpired:
        raise ValueError(".doc 转换超时（120s），请尝试上传较小的文件")

    converted = Path(out_dir) / f"{stem}.docx"
    if not converted.exists():
        raise ValueError(
            "该文件实际为旧版 .doc 格式，自动转换失败。"
            "请先用 Word 或 WPS 将文件「另存为」.docx 格式后重新上传。"
        )

    if progress_callback:
        progress_callback("格式转换完成，开始解析...", 0, -1)

    return parse_docx(str(converted), filename, progress_callback)
