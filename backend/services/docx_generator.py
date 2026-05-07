"""
DOCX 生成服务
将 Markdown 章节内容转换为格式规范的 Word 文档（.docx）
依赖：python-docx（已在 requirements.txt 中）
"""

import io
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# 公共入口
# ─────────────────────────────────────────────

def sections_to_docx(
    sections: list[dict],
    doc_title: str = "技术方案",
) -> bytes:
    """
    将章节数据列表转为 DOCX 字节流。

    sections 格式：
        [{"id": "s1", "title": "...", "level": 1, "content": "...", "done": True}, ...]

    Returns:
        bytes — DOCX 文件内容
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── 页面设置（A4，2.5 cm 边距）──────────────
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

    # ── 字体默认（宋体/仿宋正文）────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "宋体"
    style_normal.font.size = Pt(11)

    _apply_heading_styles(doc)

    # ── 文档标题 ──────────────────────────────
    title_para = doc.add_heading(doc_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 各章节 ────────────────────────────────
    for sec_data in sections:
        if not sec_data.get("done") or not (sec_data.get("content") or "").strip():
            continue

        sec_level = min(int(sec_data.get("level", 1)), 4)
        sec_title = sec_data.get("title", "")
        sec_content = sec_data.get("content", "")

        doc.add_heading(sec_title, level=sec_level)
        _parse_markdown(doc, sec_content)

    # ── 序列化 ────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# 样式设置
# ─────────────────────────────────────────────

def _apply_heading_styles(doc):
    """微调内置 Heading 样式的颜色与间距"""
    from docx.shared import Pt, RGBColor

    mapping = {
        "Heading 1": (16, "1F3864", 240, 120),
        "Heading 2": (14, "2E5496", 200, 80),
        "Heading 3": (12, "2E75B6", 160, 60),
        "Heading 4": (11, "404040", 120, 40),
    }
    for name, (size, color_hex, before, after) in mapping.items():
        try:
            st = doc.styles[name]
            st.font.size = Pt(size)
            r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
            st.font.color.rgb = RGBColor(r, g, b)
            st.paragraph_format.space_before = Pt(before / 20)
            st.paragraph_format.space_after  = Pt(after  / 20)
        except Exception:
            pass  # 样式不存在时静默忽略


# ─────────────────────────────────────────────
# Markdown 解析（行级别）
# ─────────────────────────────────────────────

def _parse_markdown(doc, md_text: str):
    """将 Markdown 文本逐行解析并插入 doc"""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        s = line.strip()

        # ── 空行 ──────────────────────────────
        if not s:
            i += 1
            continue

        # ── 标题 # ────────────────────────────
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # ── 代码块 ``` ────────────────────────
        if s.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结尾 ```
            _add_code_block(doc, "\n".join(code_lines))
            continue

        # ── 水平分隔线 ────────────────────────
        if re.match(r'^[-*_]{3,}$', s):
            i += 1
            continue

        # ── 表格 | ────────────────────────────
        if s.startswith("|") and "|" in s[1:]:
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, tbl_lines)
            continue

        # ── 无序列表 ──────────────────────────
        if re.match(r'^[-*+] ', s):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, s[2:])
            i += 1
            continue

        # ── 有序列表 ──────────────────────────
        if re.match(r'^\d+\. ', s):
            text = re.sub(r'^\d+\. ', '', s)
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, text)
            i += 1
            continue

        # ── 普通段落 ──────────────────────────
        p = doc.add_paragraph()
        _add_runs(p, s)
        i += 1


# ─────────────────────────────────────────────
# 辅助：行内格式
# ─────────────────────────────────────────────

_INLINE_PAT = re.compile(
    r'`([^`]+)`'           # `code`
    r'|\*\*([^*]+)\*\*'   # **bold**
    r'|\*([^*]+)\*'        # *italic*
    r'|([^`*]+)'           # plain
)


def _add_runs(para, text: str):
    """解析行内 Markdown（`code`、**bold**、*italic*）并添加 Run"""
    from docx.shared import Pt, RGBColor

    for m in _INLINE_PAT.finditer(text):
        code, bold, italic, plain = m.groups()
        if code:
            run = para.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        elif bold:
            run = para.add_run(bold)
            run.bold = True
        elif italic:
            run = para.add_run(italic)
            run.italic = True
        elif plain:
            para.add_run(plain)


def _add_code_block(doc, code_text: str):
    """添加代码块段落（Courier New 灰底）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)

    # 灰色背景通过段落底纹实现
    try:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        pPr.append(shd)
    except Exception:
        pass

    run = p.add_run(code_text or " ")
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def _add_table(doc, tbl_lines: list[str]):
    """将 Markdown 表格转为 DOCX 表格"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    def parse_row(line: str) -> list[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    # 过滤分隔行（:---、----）
    data_rows = [l for l in tbl_lines if not re.match(r'^[\s|:\-]+$', l)]
    if not data_rows:
        return

    headers   = parse_row(data_rows[0])
    col_count = len(headers)
    body_rows = [parse_row(r) for r in data_rows[1:]]

    table = doc.add_table(rows=1 + len(body_rows), cols=col_count)
    table.style = "Table Grid"

    # 表头（加粗 + 浅蓝底）
    hdr_row = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        try:
            from docx.oxml import OxmlElement
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "DBEAFE")
            tcPr.append(shd)
        except Exception:
            pass

    # 数据行
    for i, row_data in enumerate(body_rows):
        for j in range(min(len(row_data), col_count)):
            table.rows[i + 1].cells[j].text = row_data[j]

    doc.add_paragraph()  # 表格后空行
