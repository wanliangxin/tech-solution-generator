"""
集成测试 — 覆盖所有 REST 接口（无需真实 API Key）

运行方式：
    cd tech-solution-generator
    pip install -r requirements-dev.txt
    pytest tests/ -v
"""
import io
import pytest
from fastapi.testclient import TestClient

# conftest.py 已将 backend/ 加入 sys.path
from main import app


# ══════════════════════════════════════════════════
# Fixtures
# ══════════════════════════════════════════════════

@pytest.fixture(scope="module")
def client():
    """共用一个 TestClient 实例（module 级别，加快测试速度）"""
    with TestClient(app, raise_server_exceptions=False) as c:
        yield c


@pytest.fixture(autouse=True)
def clear_config(client):
    """每个测试前清除全局配置，避免测试间状态污染"""
    client.delete("/api/config")
    yield
    client.delete("/api/config")


def make_docx_bytes() -> bytes:
    """生成一个含标准 Heading 样式的最小 DOCX 文件（内存中构造）"""
    from docx import Document
    doc = Document()
    doc.add_heading("测试技术规范书", level=0)   # Title 样式
    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph("本项目旨在建设一套测试系统，提升工作效率。")
    doc.add_heading("1.1 建设背景", level=2)
    doc.add_paragraph("当前业务面临诸多挑战，需要引入新技术手段加以解决。")
    doc.add_heading("1.2 建设目标", level=2)
    doc.add_paragraph("构建稳定、可扩展的技术平台。")
    doc.add_heading("2. 技术架构", level=1)
    doc.add_paragraph("采用微服务架构，前后端分离，支持水平扩展。")
    doc.add_heading("3. 实施方案", level=1)
    doc.add_paragraph("分三期实施，每期约三个月，逐步上线各功能模块。")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════
# 系统
# ══════════════════════════════════════════════════

class TestSystem:
    def test_health(self, client):
        """GET /health 返回 ok"""
        r = client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"

    def test_root(self, client):
        """GET / 返回 HTML 或 JSON（取决于 frontend/index.html 是否存在）"""
        r = client.get("/")
        assert r.status_code == 200

    def test_openapi_docs(self, client):
        """GET /docs 返回 200"""
        r = client.get("/docs")
        assert r.status_code == 200


# ══════════════════════════════════════════════════
# 上传接口
# ══════════════════════════════════════════════════

class TestUpload:
    def test_upload_docx_success(self, client):
        """上传合法 DOCX → 返回结构化目录"""
        docx_bytes = make_docx_bytes()
        r = client.post(
            "/api/upload",
            files={"file": ("规范书.docx", io.BytesIO(docx_bytes),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "doc_id" in data
        assert "title" in data
        assert "sections" in data
        assert len(data["sections"]) >= 3  # 至少 3 个章节
        # 验证 section 结构
        sec = data["sections"][0]
        assert "id" in sec
        assert "level" in sec
        assert "title" in sec
        assert "content_hint" in sec

    def test_upload_invalid_format(self, client):
        """上传不支持的格式（.txt）→ 400"""
        r = client.post(
            "/api/upload",
            files={"file": ("test.txt", io.BytesIO(b"hello"), "text/plain")},
        )
        assert r.status_code == 400
        assert "格式" in r.json()["detail"]

    def test_upload_empty_file(self, client):
        """上传空文件 → 400"""
        r = client.post(
            "/api/upload",
            files={"file": ("empty.docx", io.BytesIO(b""),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert r.status_code == 400

    def test_upload_no_filename(self, client):
        """文件名为空 → 400"""
        r = client.post(
            "/api/upload",
            files={"file": ("", io.BytesIO(b"data"), "application/octet-stream")},
        )
        assert r.status_code == 400

    def test_supported_formats(self, client):
        """GET /api/upload/supported-formats 返回格式列表"""
        r = client.get("/api/upload/supported-formats")
        assert r.status_code == 200
        data = r.json()
        assert ".pdf" in data["formats"]
        assert ".docx" in data["formats"]
        assert data["max_size_mb"] == 50


# ══════════════════════════════════════════════════
# 配置接口
# ══════════════════════════════════════════════════

class TestConfig:
    def test_get_config_not_configured(self, client):
        """未配置时 GET /api/config → configured=false"""
        r = client.get("/api/config")
        assert r.status_code == 200
        assert r.json()["configured"] is False

    def test_save_config_openai(self, client):
        """POST /api/config 保存 openai 配置"""
        r = client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-test1234567890",
            "base_url": "https://api.openai.com/v1",
            "model": "gpt-4o",
        })
        assert r.status_code == 200
        data = r.json()
        assert data["success"] is True
        cfg = data["config"]
        assert cfg["provider"] == "openai"
        assert cfg["model"] == "gpt-4o"
        # Key 应脱敏（不含完整原始 key）
        assert "test1234567890" not in cfg["api_key_masked"]

    def test_save_config_claude(self, client):
        """POST /api/config 保存 claude 配置"""
        r = client.post("/api/config", json={
            "provider": "claude",
            "api_key": "sk-ant-test12345678",
        })
        assert r.status_code == 200
        data = r.json()
        assert data["config"]["provider"] == "claude"
        # 未填 model 时应使用默认值
        assert data["config"]["model"] != ""

    def test_save_config_provider_case_insensitive(self, client):
        """provider 大小写不敏感"""
        r = client.post("/api/config", json={
            "provider": "OpenAI",
            "api_key": "sk-test1234567890",
        })
        assert r.status_code == 200

    def test_save_config_invalid_provider(self, client):
        """非法 provider → 422"""
        r = client.post("/api/config", json={
            "provider": "gemini",
            "api_key": "sk-test1234567890",
        })
        assert r.status_code == 422

    def test_save_config_empty_key(self, client):
        """api_key 为空 → 422"""
        r = client.post("/api/config", json={
            "provider": "openai",
            "api_key": "",
        })
        assert r.status_code == 422

    def test_save_config_short_key(self, client):
        """api_key 过短 → 422"""
        r = client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-123",
        })
        assert r.status_code == 422

    def test_get_config_after_save(self, client):
        """保存后 GET /api/config → configured=true，key 脱敏"""
        client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-test1234567890abcdef",
            "model": "gpt-4o-mini",
        })
        r = client.get("/api/config")
        assert r.status_code == 200
        data = r.json()
        assert data["configured"] is True
        cfg = data["config"]
        assert cfg["provider"] == "openai"
        assert cfg["model"] == "gpt-4o-mini"
        assert "abcdef" not in cfg["api_key_masked"]   # 中间被遮盖

    def test_delete_config(self, client):
        """DELETE /api/config → 清除配置"""
        client.post("/api/config", json={"provider": "openai", "api_key": "sk-test1234567890"})
        r = client.delete("/api/config")
        assert r.status_code == 200
        assert r.json()["success"] is True
        # 验证已清除
        r2 = client.get("/api/config")
        assert r2.json()["configured"] is False

    def test_check_config_not_configured(self, client):
        """未配置时 GET /api/config/check → 400"""
        r = client.get("/api/config/check")
        assert r.status_code == 400

    def test_list_providers(self, client):
        """GET /api/config/providers 返回 openai 和 claude"""
        r = client.get("/api/config/providers")
        assert r.status_code == 200
        providers = {p["id"] for p in r.json()["providers"]}
        assert "openai" in providers
        assert "claude" in providers


# ══════════════════════════════════════════════════
# 生成接口
# ══════════════════════════════════════════════════

SAMPLE_SECTIONS = [
    {"id": "s1", "title": "1. 项目概述", "content": "项目背景与目标说明"},
    {"id": "s2", "title": "2. 技术架构", "content": "微服务架构设计"},
]

class TestGenerate:
    def _save_config(self, client):
        """辅助：保存一个假 key（不验证有效性）"""
        client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-fake-key-for-testing-only",
            "base_url": "https://api.openai.com/v1",
            "model": "gpt-4o",
        })

    def test_start_without_config(self, client):
        """未配置 → POST /api/generate/start 返回 400"""
        r = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        assert r.status_code == 400

    def test_start_empty_sections(self, client):
        """sections 为空列表 → 422"""
        self._save_config(client)
        r = client.post("/api/generate/start", json={"sections": []})
        assert r.status_code == 422

    def test_start_too_many_sections(self, client):
        """sections 超过 50 个 → 422"""
        self._save_config(client)
        sections = [{"id": f"s{i}", "title": f"章节{i}", "content": ""}
                    for i in range(51)]
        r = client.post("/api/generate/start", json={"sections": sections})
        assert r.status_code == 422

    def test_start_success(self, client):
        """配置完成后 POST /api/generate/start → 返回 task_id"""
        self._save_config(client)
        r = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        assert r.status_code == 200, r.text
        data = r.json()
        assert "task_id" in data
        assert "stream_url" in data
        assert "status_url" in data
        assert data["total_sections"] == 2
        return data["task_id"]

    def test_status_not_found(self, client):
        """未知 task_id → GET /api/generate/status/{id} 返回 404"""
        r = client.get("/api/generate/status/nonexistent-task-id")
        assert r.status_code == 404

    def test_status_after_start(self, client):
        """启动后立即查询状态 → 有效的状态字段"""
        self._save_config(client)
        r_start = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        task_id = r_start.json()["task_id"]

        r_status = client.get(f"/api/generate/status/{task_id}")
        assert r_status.status_code == 200
        data = r_status.json()
        assert data["task_id"] == task_id
        assert data["status"] in ("pending", "running", "completed", "error", "cancelled")
        assert "progress" in data
        assert "total_sections" in data
        assert data["total_sections"] == 2

    def test_cancel_not_found(self, client):
        """取消不存在的任务 → 404"""
        r = client.delete("/api/generate/nonexistent-task-id")
        assert r.status_code == 404

    def test_cancel_pending_task(self, client):
        """启动任务后立即取消（pending 状态）→ 200"""
        self._save_config(client)
        r_start = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        task_id = r_start.json()["task_id"]
        # 立即取消（可能是 pending 或 running）
        r_cancel = client.delete(f"/api/generate/{task_id}")
        assert r_cancel.status_code == 200


# ══════════════════════════════════════════════════
# 下载接口
# ══════════════════════════════════════════════════

class TestDownload:
    def test_download_not_found(self, client):
        """未知 task_id → 404"""
        r = client.get("/api/download/nonexistent-task-id")
        assert r.status_code == 404

    def test_download_pending_task(self, client):
        """pending 任务 → 425（内容未生成）"""
        client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-fake-key-for-testing-only",
        })
        r_start = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        task_id = r_start.json()["task_id"]
        # 立即取消，使任务停在 pending/running（未产生内容）
        client.delete(f"/api/generate/{task_id}")

        r_dl = client.get(f"/api/download/{task_id}")
        # pending 返回 425，running 但无内容返回 404
        assert r_dl.status_code in (404, 425)

    def test_download_json_not_found(self, client):
        """JSON 下载：未知 task_id → 404"""
        r = client.get("/api/download/nonexistent-task-id/json")
        assert r.status_code == 404

    def test_download_json_pending_task(self, client):
        """JSON 下载：pending 任务 → 返回空 sections 列表"""
        client.post("/api/config", json={
            "provider": "openai",
            "api_key": "sk-fake-key-for-testing-only",
        })
        r_start = client.post("/api/generate/start", json={"sections": SAMPLE_SECTIONS})
        task_id = r_start.json()["task_id"]

        r_dl = client.get(f"/api/download/{task_id}/json")
        assert r_dl.status_code == 200
        data = r_dl.json()
        assert data["task_id"] == task_id
        assert "sections" in data
        assert len(data["sections"]) == 2
